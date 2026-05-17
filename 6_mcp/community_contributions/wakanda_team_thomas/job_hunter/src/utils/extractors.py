"""Text extraction utilities for PDF and DOCX files."""

from pathlib import Path
from typing import Union

from pypdf import PdfReader
from docx import Document


MAX_FILE_SIZE_BYTES = 1 * 1024 * 1024  # 1 MB


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def validate_file_size(path: Path) -> None:
    """
    Validate that file size is within the allowed limit.
    
    Raises:
        ExtractionError: If file exceeds MAX_FILE_SIZE_BYTES.
    """
    file_size = path.stat().st_size
    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        raise ExtractionError(
            f"File size ({size_mb:.2f} MB) exceeds maximum allowed ({max_mb:.0f} MB)"
        )


def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file.
        
    Returns:
        Extracted text as a single string.
        
    Raises:
        ExtractionError: If the file cannot be read or parsed.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise ExtractionError(f"File not found: {path}")
    
    if path.suffix.lower() != ".pdf":
        raise ExtractionError(f"Expected PDF file, got: {path.suffix}")
    
    validate_file_size(path)
    
    try:
        reader = PdfReader(str(path))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts).strip()
        
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from PDF: {e}") from e


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file.
        
    Returns:
        Extracted text as a single string.
        
    Raises:
        ExtractionError: If the file cannot be read or parsed.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise ExtractionError(f"File not found: {path}")
    
    if path.suffix.lower() != ".docx":
        raise ExtractionError(f"Expected DOCX file, got: {path.suffix}")
    
    validate_file_size(path)
    
    try:
        doc = Document(str(path))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts).strip()
        
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from DOCX: {e}") from e


def extract_text(file_path: Union[str, Path]) -> str:
    """
    Extract text from a file based on its extension.
    
    Supports PDF and DOCX formats.
    
    Args:
        file_path: Path to the file.
        
    Returns:
        Extracted text as a single string.
        
    Raises:
        ExtractionError: If the file format is unsupported or extraction fails.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    elif suffix == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ExtractionError(f"Unsupported file format: {suffix}. Supported: .pdf, .docx")


def get_supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return [".pdf", ".docx"]


def is_supported_file(file_path: Union[str, Path]) -> bool:
    """Check if a file has a supported extension."""
    path = Path(file_path)
    return path.suffix.lower() in get_supported_extensions()
