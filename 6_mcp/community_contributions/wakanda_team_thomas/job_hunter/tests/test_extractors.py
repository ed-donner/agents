"""Tests for text extraction utilities."""

from pathlib import Path

import pytest
from pypdf import PdfWriter
from docx import Document

from src.utils.extractors import (
    MAX_FILE_SIZE_BYTES,
    ExtractionError,
    extract_text,
    extract_text_from_pdf,
    extract_text_from_docx,
    get_supported_extensions,
    is_supported_file,
    validate_file_size,
)


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    """Create a sample PDF file for testing."""
    pdf_path = tmp_path / "resume.pdf"
    
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    
    with open(pdf_path, "wb") as f:
        writer.write(f)
    
    return pdf_path


@pytest.fixture
def sample_pdf_with_text(tmp_path) -> Path:
    """Create a PDF with actual text content."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    pdf_path = tmp_path / "resume_with_text.pdf"
    
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "John Doe")
    c.drawString(100, 730, "Software Engineer")
    c.drawString(100, 710, "john.doe@email.com")
    c.drawString(100, 680, "Skills: Python, JavaScript, AWS")
    c.save()
    
    return pdf_path


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """Create a sample DOCX file for testing."""
    docx_path = tmp_path / "resume.docx"
    
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("Senior Developer")
    doc.add_paragraph("jane.smith@email.com")
    doc.add_paragraph("")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, React, Docker, Kubernetes")
    doc.add_paragraph("")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Tech Corp - Lead Engineer - 2020-Present")
    
    doc.save(docx_path)
    
    return docx_path


@pytest.fixture
def sample_docx_with_table(tmp_path) -> Path:
    """Create a DOCX file with a table."""
    docx_path = tmp_path / "resume_table.docx"
    
    doc = Document()
    doc.add_paragraph("Resume")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"
    
    doc.save(docx_path)
    
    return docx_path


class TestPdfExtraction:
    """Tests for PDF extraction."""

    def test_extract_from_nonexistent_file(self):
        """Test extraction from non-existent file raises error."""
        with pytest.raises(ExtractionError, match="File not found"):
            extract_text_from_pdf("/nonexistent/file.pdf")

    def test_extract_from_wrong_extension(self, tmp_path):
        """Test extraction from wrong file type raises error."""
        txt_file = tmp_path / "file.txt"
        txt_file.write_text("content")
        
        with pytest.raises(ExtractionError, match="Expected PDF file"):
            extract_text_from_pdf(txt_file)

    def test_extract_from_blank_pdf(self, sample_pdf):
        """Test extraction from blank PDF returns empty string."""
        text = extract_text_from_pdf(sample_pdf)
        assert text == ""

    def test_extract_pdf_with_text(self, sample_pdf_with_text):
        """Test extraction from PDF with text content."""
        text = extract_text_from_pdf(sample_pdf_with_text)
        
        assert "John Doe" in text
        assert "Software Engineer" in text
        assert "Python" in text


class TestDocxExtraction:
    """Tests for DOCX extraction."""

    def test_extract_from_nonexistent_file(self):
        """Test extraction from non-existent file raises error."""
        with pytest.raises(ExtractionError, match="File not found"):
            extract_text_from_docx("/nonexistent/file.docx")

    def test_extract_from_wrong_extension(self, tmp_path):
        """Test extraction from wrong file type raises error."""
        txt_file = tmp_path / "file.txt"
        txt_file.write_text("content")
        
        with pytest.raises(ExtractionError, match="Expected DOCX file"):
            extract_text_from_docx(txt_file)

    def test_extract_from_docx(self, sample_docx):
        """Test extraction from DOCX file."""
        text = extract_text_from_docx(sample_docx)
        
        assert "Jane Smith" in text
        assert "Senior Developer" in text
        assert "Python" in text
        assert "Tech Corp" in text

    def test_extract_docx_with_table(self, sample_docx_with_table):
        """Test extraction from DOCX with table."""
        text = extract_text_from_docx(sample_docx_with_table)
        
        assert "Python" in text
        assert "Expert" in text


class TestGenericExtraction:
    """Tests for generic extract_text function."""

    def test_extract_pdf(self, sample_pdf_with_text):
        """Test generic extraction detects PDF."""
        text = extract_text(sample_pdf_with_text)
        assert "John Doe" in text

    def test_extract_docx(self, sample_docx):
        """Test generic extraction detects DOCX."""
        text = extract_text(sample_docx)
        assert "Jane Smith" in text

    def test_unsupported_format(self, tmp_path):
        """Test unsupported format raises error."""
        txt_file = tmp_path / "file.txt"
        txt_file.write_text("content")
        
        with pytest.raises(ExtractionError, match="Unsupported file format"):
            extract_text(txt_file)


class TestHelperFunctions:
    """Tests for helper functions."""

    def test_supported_extensions(self):
        """Test get_supported_extensions returns correct list."""
        extensions = get_supported_extensions()
        assert ".pdf" in extensions
        assert ".docx" in extensions
        assert len(extensions) == 2

    def test_is_supported_pdf(self, tmp_path):
        """Test is_supported_file for PDF."""
        pdf_path = tmp_path / "file.pdf"
        assert is_supported_file(pdf_path) is True

    def test_is_supported_docx(self, tmp_path):
        """Test is_supported_file for DOCX."""
        docx_path = tmp_path / "file.docx"
        assert is_supported_file(docx_path) is True

    def test_is_not_supported_txt(self, tmp_path):
        """Test is_supported_file for unsupported format."""
        txt_path = tmp_path / "file.txt"
        assert is_supported_file(txt_path) is False

    def test_is_supported_case_insensitive(self, tmp_path):
        """Test is_supported_file is case insensitive."""
        pdf_path = tmp_path / "file.PDF"
        assert is_supported_file(pdf_path) is True


class TestFileSizeValidation:
    """Tests for file size validation."""

    def test_max_file_size_is_one_mb(self):
        """Test MAX_FILE_SIZE_BYTES is 1 MB."""
        assert MAX_FILE_SIZE_BYTES == 1 * 1024 * 1024

    def test_file_within_limit(self, tmp_path):
        """Test file within size limit passes validation."""
        small_file = tmp_path / "small.txt"
        small_file.write_bytes(b"x" * 1000)
        
        validate_file_size(small_file)

    def test_file_exceeds_limit(self, tmp_path):
        """Test file exceeding size limit raises error."""
        large_file = tmp_path / "large.txt"
        large_file.write_bytes(b"x" * (MAX_FILE_SIZE_BYTES + 1))
        
        with pytest.raises(ExtractionError, match="exceeds maximum"):
            validate_file_size(large_file)

    def test_pdf_extraction_rejects_large_file(self, tmp_path):
        """Test PDF extraction rejects files over 1 MB."""
        large_pdf = tmp_path / "large.pdf"
        large_pdf.write_bytes(b"x" * (MAX_FILE_SIZE_BYTES + 1))
        
        with pytest.raises(ExtractionError, match="exceeds maximum"):
            extract_text_from_pdf(large_pdf)

    def test_docx_extraction_rejects_large_file(self, tmp_path):
        """Test DOCX extraction rejects files over 1 MB."""
        large_docx = tmp_path / "large.docx"
        large_docx.write_bytes(b"x" * (MAX_FILE_SIZE_BYTES + 1))
        
        with pytest.raises(ExtractionError, match="exceeds maximum"):
            extract_text_from_docx(large_docx)
