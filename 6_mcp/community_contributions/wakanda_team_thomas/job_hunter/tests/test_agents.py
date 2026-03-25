"""Tests for OpenAI Agents."""

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from src.agent_workers.resume_parser import (
    ParsedResume,
    ParsedSkill,
    ParsedExperience,
    ParsedEducation,
    _extract_resume_text_impl,
    extract_resume_text,
    resume_parser_agent,
)


class TestParsedResumeSchema:
    """Test ParsedResume Pydantic schema."""

    def test_minimal_resume(self):
        """Test creating a minimal parsed resume."""
        resume = ParsedResume(name="John Doe")
        
        assert resume.name == "John Doe"
        assert resume.email is None
        assert resume.skills == []
        assert resume.experience == []

    def test_full_resume(self):
        """Test creating a fully populated parsed resume."""
        resume = ParsedResume(
            name="Jane Smith",
            email="jane@example.com",
            phone="+1-555-123-4567",
            location="San Francisco, CA",
            summary="Senior software engineer with 10 years experience",
            skills=[
                ParsedSkill(name="Python", level="expert", years=8),
                ParsedSkill(name="AWS", level="advanced", years=5),
            ],
            experience=[
                ParsedExperience(
                    company="Tech Corp",
                    title="Senior Engineer",
                    start_date="2020-01",
                    end_date="Present",
                    description="Led backend development team",
                ),
            ],
            education=[
                ParsedEducation(
                    institution="MIT",
                    degree="Bachelor's",
                    field="Computer Science",
                    graduation_date="2014",
                ),
            ],
            certifications=["AWS Solutions Architect"],
            languages=["English", "Spanish"],
            keywords=["python", "aws", "backend", "microservices"],
        )
        
        assert resume.name == "Jane Smith"
        assert len(resume.skills) == 2
        assert resume.skills[0].name == "Python"
        assert len(resume.experience) == 1
        assert resume.experience[0].company == "Tech Corp"


class TestParsedSkillSchema:
    """Test ParsedSkill schema."""

    def test_skill_with_all_fields(self):
        """Test skill with all optional fields."""
        skill = ParsedSkill(name="Python", level="expert", years=5)
        
        assert skill.name == "Python"
        assert skill.level == "expert"
        assert skill.years == 5

    def test_skill_minimal(self):
        """Test skill with only required name."""
        skill = ParsedSkill(name="Docker")
        
        assert skill.name == "Docker"
        assert skill.level is None
        assert skill.years is None


class TestParsedExperienceSchema:
    """Test ParsedExperience schema."""

    def test_experience_full(self):
        """Test experience with all fields."""
        exp = ParsedExperience(
            company="Startup Inc",
            title="CTO",
            start_date="2019-06",
            end_date="2023-12",
            description="Built engineering team from scratch",
            location="Remote",
        )
        
        assert exp.company == "Startup Inc"
        assert exp.title == "CTO"
        assert exp.location == "Remote"

    def test_experience_minimal(self):
        """Test experience with only required fields."""
        exp = ParsedExperience(company="Company", title="Engineer")
        
        assert exp.company == "Company"
        assert exp.title == "Engineer"
        assert exp.start_date is None


class TestParsedEducationSchema:
    """Test ParsedEducation schema."""

    def test_education_full(self):
        """Test education with all fields."""
        edu = ParsedEducation(
            institution="Stanford University",
            degree="Master's",
            field="Machine Learning",
            graduation_date="2018-06",
        )
        
        assert edu.institution == "Stanford University"
        assert edu.degree == "Master's"
        assert edu.field == "Machine Learning"

    def test_education_minimal(self):
        """Test education with only institution."""
        edu = ParsedEducation(institution="Community College")
        
        assert edu.institution == "Community College"
        assert edu.degree is None


class TestExtractResumeTool:
    """Test the extract_resume_text function tool."""

    def test_unsupported_file_format(self, tmp_path):
        """Test that unsupported formats return error message."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text("Some content")
        
        result = _extract_resume_text_impl(str(txt_file))
        
        assert "Error" in result
        assert "Unsupported" in result

    def test_nonexistent_file(self):
        """Test that nonexistent file returns error."""
        result = _extract_resume_text_impl("/nonexistent/resume.pdf")
        
        assert "Error" in result

    def test_extract_from_docx(self, tmp_path):
        """Test extraction from a DOCX file."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Software Engineer")
        doc.add_paragraph("Python, JavaScript, AWS")
        
        docx_path = tmp_path / "resume.docx"
        doc.save(str(docx_path))
        
        result = _extract_resume_text_impl(str(docx_path))
        
        assert "John Doe" in result
        assert "Software Engineer" in result
        assert "Python" in result


class TestResumeParserAgent:
    """Test the resume parser agent configuration."""

    def test_agent_has_correct_name(self):
        """Test agent has expected name."""
        assert resume_parser_agent.name == "ResumeParserAgent"

    def test_agent_has_tools(self):
        """Test agent has the extract tool."""
        tool_names = [t.name for t in resume_parser_agent.tools]
        assert "extract_resume_text" in tool_names

    def test_agent_has_output_type(self):
        """Test agent has structured output type."""
        assert resume_parser_agent.output_type == ParsedResume
